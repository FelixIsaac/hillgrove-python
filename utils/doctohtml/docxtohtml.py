from os import listdir
from os.path import isfile, join
import docx

DOC_DIR = './docs/'
OUT_DIR = './out/'
docs = [f for f in listdir(DOC_DIR) if isfile(join(DOC_DIR, f))]

def write_file(html, doc_name):
    f = open(OUT_DIR + doc_name.split('.')[0] + '.tsx', 'a')
    f.write(html)
    f.close()

def generate(doc):
    segment_title = doc.paragraphs[0]
    print(segment_title.style)

    def text(paragraph):
        text = paragraph.text.split(' ')
        chunks, chunk_size = len(text), len(text) // 3

        if chunk_size == 0: return

        formattedText = [ ' '.join(text[i : i + chunk_size]) for i in range(0, chunks, chunk_size) ]

        html = '''
            <Text>
            \t{}
            <Text>'''

        return html.format('\n\t\t\t\t'.join(formattedText))

    def format_chain(paragraph):
        formatted = text(paragraph)

        return formatted

    formatted_paragraphs = ''.join(filter(lambda x: x, list(map(format_chain, doc.paragraphs))))
    
    result = '''
import React from 'react';
import {{ Heading, Text, Image, Center, Box }} from '@chakra-ui/react';    

const SegmentManager = () => {{
    return (
        <Box as="article" fontSize="18px">
            <Heading as="h2">{}</Heading>{}
        </Box>
    )
}}

export default SegmentManager;
'''.format(segment_title.text, formatted_paragraphs)

    return result

for doc in docs:
    result = generate(docx.Document(DOC_DIR + doc))
    write_file(result, doc)